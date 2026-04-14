"""
The Orchestrator – State Manager & ASI:One Gateway
====================================================
Run in its own terminal (start AFTER the three sub-agents):

    python orchestrator.py

This is the ONLY agent that faces the outside world. It does four things:
  1. Catches incoming messages from ASI:One via the Agent Chat Protocol.
  2. Strips PDF text using pdfplumber.
  3. Chains Extractor → Scriptwriter → VoiceStudio with send_and_receive.
  4. Delivers the final MP3 back to the user.

It also exposes a REST POST /process endpoint for local testing.

Sub-agent addresses are read from environment variables. Run
`python get_addresses.py` first to print the addresses you need to set.
"""

import asyncio
import base64
import io
import json
import os
import re
import sys
import threading
from datetime import datetime, timezone
from functools import partial
from http.server import HTTPServer, SimpleHTTPRequestHandler
from pathlib import Path
from typing import Optional
from uuid import UUID, uuid4

import requests  # type: ignore[import-untyped]
import stripe as _stripe_lib
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from openai import AsyncOpenAI
import pdfplumber
from uagents import Agent, Context, Protocol

from dotenv import load_dotenv

from uagents_core.contrib.protocols.chat import (
    ChatAcknowledgement,
    ChatMessage,
    EndSessionContent,
    Resource,
    ResourceContent,
    TextContent,
    chat_protocol_spec,
)

from uagents_core.contrib.protocols.payment import (
    CommitPayment,
    CompletePayment,
    Funds,
    RejectPayment,
    RequestPayment,
    payment_protocol_spec,
)

from schemas import (
    AudioResponse,
    ContextInjection,
    DebateResponse,
    DebateTurn,
    ExtractRequest,
    PipelineRequest,
    PipelineResponse,
    PodcastScript,
    ResearchInsights,
)

load_dotenv()

# Windows PowerShell defaults to cp1252 which can't encode arrows/checkmarks.
# Force UTF-8 on stdout/stderr so logger never crashes on Unicode characters.
if sys.stdout and hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(
        sys.stdout.buffer, encoding="utf-8", errors="replace", line_buffering=True
    )
if sys.stderr and hasattr(sys.stderr, "buffer"):
    sys.stderr = io.TextIOWrapper(
        sys.stderr.buffer, encoding="utf-8", errors="replace", line_buffering=True
    )

# ── Sub-agent addresses (set via env vars — run get_addresses.py first) ───────

EXTRACTOR_ADDRESS = os.getenv("EXTRACTOR_ADDRESS", "")
SCRIPTWRITER_ADDRESS = os.getenv("SCRIPTWRITER_ADDRESS", "")
VOICE_STUDIO_ADDRESS = os.getenv("VOICE_STUDIO_ADDRESS", "")
HOST_A_ADDRESS = os.getenv("HOST_A_ADDRESS", "")
HOST_B_ADDRESS = os.getenv("HOST_B_ADDRESS", "")

OUTPUT_DIR = Path(os.getenv("OUTPUT_DIR", "output"))
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

AUDIO_SERVER_PORT = int(os.getenv("AUDIO_SERVER_PORT", "8080"))
AUDIO_SERVER_HOST = os.getenv("AUDIO_SERVER_HOST", "localhost")

# ── Live debate trigger keywords & storage key ────────────────────────────────
_LIVE_DEBATE_TRIGGERS = {
    "live debate",
    "start debate",
    "watch debate",
    "debate now",
    "watch them",
    "let them talk",
    "replay",
    "show debate",
    "debate",
    "continue debate",
    "continue",
}
_LAST_SESSION_KEY = "last_session_context"
_ACTIVE_DEBATE_KEY = "active_debate_session"
_DEBATE_ACCUM_KEY = "debate_transcript_accum"
_DEBATE_COOLDOWN_KEY = (
    "debate_last_started_ts"  # Unix timestamp — debounces duplicate triggers
)
_DEBATE_MSG_ID_KEY = (
    "debate_bubble_msg_id"  # stable UUID so all updates hit the same bubble
)
_PERSONALITIES_KEY = "host_personalities"
_PAID_SESSIONS_KEY = "paid_podcast_sessions"  # JSON list of paid session IDs
_PENDING_PAYMENTS_KEY = (
    "pending_stripe_payments"  # JSON dict keyed by checkout_session_id
)

_DEBATE_COOLDOWN_SECS = 30  # ignore duplicate "debate" triggers within this window
_ACTIVE_DEBATE_TTL = 300  # auto-clear stale active debate after 5 minutes
_ACTIVE_DEBATE_TS_KEY = "active_debate_started_ts"  # Unix timestamp of debate start
_SEEN_MSG_IDS_KEY = "seen_msg_ids"  # JSON list of recently processed msg_ids

# ── Personality presets ───────────────────────────────────────────────────────
# Each entry: (display_name, system_prompt_hint)
_PERSONALITY_PRESETS: dict[str, dict[str, tuple[str, str]]] = {
    "host_a": {
        "1": (
            "Classic Skeptic",
            "You demand hard empirical evidence. You always ask 'what's the sample size?' "
            "and 'was this peer-reviewed?'. You challenge methodology relentlessly but fairly. "
            "You are NOT dismissive — you are intellectually rigorous.",
        ),
        "2": (
            "Investigative Journalist",
            "You follow the money. You ask 'who funded this research?', probe conflicts of "
            "interest, and surface what the paper conveniently omits. You are the audience's "
            "advocate — sharp, tenacious, and never satisfied with PR-speak.",
        ),
        "3": (
            "Academic Critic",
            "You are obsessed with methodological rigor and the replication crisis. You demand "
            "statistical significance thresholds, effect sizes, confidence intervals, and "
            "independent replication before accepting any finding as credible.",
        ),
        "4": (
            "Industry Veteran",
            "You have seen every hype cycle — dot-com, blockchain, metaverse. You compare new "
            "claims to past failed promises, demand real-world deployment numbers over lab "
            "results, and are deeply sceptical of trend reports and analyst projections.",
        ),
    },
    "host_b": {
        "1": (
            "Researcher",
            "You cite exact data points from the study. You defend findings with measured "
            "confidence, acknowledge limitations honestly, and never over-claim. You often "
            "say 'the data shows' and back it up with a specific number.",
        ),
        "2": (
            "Industry Insider",
            "You speak in terms of ROI, market share, and enterprise adoption curves. "
            "You name companies already shipping these solutions, focus on practical value "
            "delivered, and translate academic findings into business impact.",
        ),
        "3": (
            "Futurist",
            "You connect the paper's findings to sweeping technological and societal shifts. "
            "You extrapolate decades ahead, frame every result as part of a larger paradigm "
            "change, and are genuinely excited about what comes next.",
        ),
        "4": (
            "Enthusiastic Teacher",
            "You use vivid analogies to make complex concepts click for anyone. You celebrate "
            "the 'aha moment', make the audience feel smart, and genuinely love helping people "
            "understand difficult ideas through relatable storytelling.",
        ),
    },
}

_PERSONALITY_SET_RE = re.compile(
    r"\bA\s*:\s*([1-4])\b.*?\bB\s*:\s*([1-4])\b", re.IGNORECASE
)
_PERSONALITY_TRIGGERS = {
    "customize",
    "set personality",
    "change personality",
    "set hosts",
    "host style",
    "personalities",
    "change hosts",
}

_PERSONALITY_MENU = (
    "## 🎭 Customize Host Personalities\n\n"
    "Reply with **A:[1-4] B:[1-4]** to set host styles "
    "*(e.g. `A:2 B:3`)*.\n\n"
    "**🎤 @skeptic-agent — the challenger**\n"
    "1. **Classic Skeptic** — demands evidence, asks for sample sizes\n"
    "2. **Investigative Journalist** — follows the money, probes conflicts of interest\n"
    "3. **Academic Critic** — obsessed with methodology, replication, and stats rigor\n"
    "4. **Industry Veteran** — has seen every hype cycle, compares to past failures\n\n"
    "**🎓 @expert-agent — the defender**\n"
    "1. **Researcher** — cites exact data points, acknowledges limitations\n"
    "2. **Industry Insider** — ROI focus, names companies, practical value\n"
    "3. **Futurist** — connects findings to paradigm shifts, extrapolates ahead\n"
    "4. **Enthusiastic Teacher** — vivid analogies, makes complex ideas accessible\n\n"
    "*Current default: A:1 (Classic Skeptic) · B:1 (Researcher)*"
)


def _chat(text: str) -> ChatMessage:
    """Create a ChatMessage that closes the current ASI:One bubble immediately.

    Without EndSessionContent, ASI:One keeps the response bubble 'open' and
    every subsequent send from this agent simply updates (overwrites) that same
    bubble.  Appending EndSessionContent seals the bubble so the next send
    opens a fresh one — giving each pipeline step and each debate turn its own
    separate chat bubble.
    """
    return ChatMessage(
        timestamp=datetime.now(timezone.utc),
        msg_id=uuid4(),
        content=[
            TextContent(type="text", text=text),
            EndSessionContent(type="end-session"),
        ],
    )


def _chat_open(text: str, msg_id: UUID) -> ChatMessage:
    """Create a ChatMessage that keeps the bubble OPEN for future updates.

    By reusing the same ``msg_id`` and omitting EndSessionContent, ASI:One
    will overwrite the existing bubble in-place rather than creating a new one.
    """
    return ChatMessage(
        timestamp=datetime.now(timezone.utc),
        msg_id=msg_id,
        content=[TextContent(type="text", text=text)],
    )


def _chat_close(text: str, msg_id: UUID) -> ChatMessage:
    """Create a ChatMessage that updates the bubble one last time and seals it."""
    return ChatMessage(
        timestamp=datetime.now(timezone.utc),
        msg_id=msg_id,
        content=[
            TextContent(type="text", text=text),
            EndSessionContent(type="end-session"),
        ],
    )


def _start_audio_server() -> None:
    """Serve OUTPUT_DIR over HTTP so MP3 links are clickable and browser-playable."""
    handler = partial(SimpleHTTPRequestHandler, directory=str(OUTPUT_DIR))
    server = HTTPServer(("0.0.0.0", AUDIO_SERVER_PORT), handler)
    server.serve_forever()


threading.Thread(target=_start_audio_server, daemon=True, name="audio-server").start()


def _audio_url(audio_path: str) -> str:
    filename = Path(audio_path).name
    return f"http://{AUDIO_SERVER_HOST}:{AUDIO_SERVER_PORT}/{filename}"


# ── PDF helpers ───────────────────────────────────────────────────────────────


def _pdf_bytes_to_text(pdf_bytes: bytes) -> str:
    lines = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                lines.append(t)
    return "\n".join(lines)


def _pdf_path_to_text(path: str) -> str:
    lines = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                lines.append(t)
    return "\n".join(lines)


def _resource_to_text(resource: Resource) -> str:
    """Fetch PDF bytes from a ResourceContent URI and extract text.

    Handles three URI schemes emitted by ASI:One:
      - https://...          → download with requests
      - data:...;base64,...  → decode inline
      - everything else      → treat as a local file path
    """
    uri = resource.uri
    if uri.startswith("http://") or uri.startswith("https://"):
        resp = requests.get(uri, timeout=30)
        resp.raise_for_status()
        pdf_bytes = resp.content
    elif uri.startswith("data:"):
        _, encoded = uri.split(",", 1)
        pdf_bytes = base64.b64decode(encoded)
    else:
        pdf_bytes = Path(uri).read_bytes()
    return _pdf_bytes_to_text(pdf_bytes)


# ── LLM clients ───────────────────────────────────────────────────────────────
# All text generation → ASI:One (stays within the Fetch.ai ecosystem)

_asi = AsyncOpenAI(
    api_key=os.getenv("ASI1_API_KEY", ""),
    base_url="https://api.asi1.ai/v1",
)

_ASI_MODEL = os.getenv("ASI1_MODEL", "asi1-mini")
_DOCX_MODEL = _ASI_MODEL


def _docx_slug(title: str) -> str:
    return re.sub(r"[^a-zA-Z0-9]+", "_", title)[:50]


def _set_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def _add_speaker_block(
    doc: Document, speaker_label: str, text: str, color: RGBColor
) -> None:
    p = doc.add_paragraph()
    label = p.add_run(f"{speaker_label}  ")
    label.bold = True
    label.font.color.rgb = color
    label.font.size = Pt(10)
    body = p.add_run(text)
    body.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(6)


async def _build_docx(
    audio_path: str,
    script: "PodcastScript",
    insights: "ResearchInsights",
) -> str:
    """Generate an extended director's-cut script via LLM and write a DOCX.

    The document contains:
      1. Cover metadata (title, date, stats)
      2. Episode Brief  (thesis · key metrics · controversy)
      3. Podcast Script — the lines that were voiced
      4. Extended Script — ~25 additional exchanges, director's cut
    Returns the saved .docx path.
    """
    COLOR_A = RGBColor(0x0D, 0x47, 0xA1)  # deep blue  – Host A (Skeptic)
    COLOR_B = RGBColor(0x1B, 0x5E, 0x20)  # deep green – Host B (Expert)
    COLOR_HDR = RGBColor(0x1A, 0x1A, 0x2E)  # near-black

    # ── Generate extended dialogue via LLM ─────────────────────────────────
    extend_prompt = (
        "You are a podcast showrunner writing a director's-cut script.\n"
        "Below is a short voiced podcast debate. Extend it into a longer written version "
        "(25–30 more exchanges) that goes deeper: explore edge cases, counter-arguments, "
        "analogies, and open questions the paper didn't answer. "
        "Keep the same two hosts: HostA (skeptic, demands proof) and HostB (expert, defends with numbers).\n\n"
        f"CONTEXT\n"
        f"Topic: {script.topic_title}\n"
        f"Thesis: {insights.core_thesis}\n"
        f"Key metrics: {', '.join(insights.key_metrics)}\n"
        f"Controversy: {insights.controversial_point}\n\n"
        "VOICED SCRIPT (do not repeat verbatim, but continue naturally from it):\n"
        + "\n".join(
            f"{'HostA' if line.speaker == 'HostA' else 'HostB'}: {line.text}"
            for line in script.lines
        )
        + "\n\n"
        "Return ONLY a JSON array of objects with keys 'speaker' ('HostA'|'HostB') and 'text'. "
        "No markdown, no extra keys."
    )

    resp = await _asi.chat.completions.create(
        model=_ASI_MODEL,
        messages=[
            {
                "role": "system",
                "content": "You are a podcast showrunner. Return only valid JSON.",
            },
            {"role": "user", "content": extend_prompt},
        ],
        temperature=0.85,
    )
    raw = resp.choices[0].message.content or "{}"
    try:
        parsed = json.loads(raw)
        # model may wrap in a key
        if isinstance(parsed, dict):
            for v in parsed.values():
                if isinstance(v, list):
                    parsed = v
                    break
        extended_lines = parsed if isinstance(parsed, list) else []
    except Exception:
        extended_lines = []

    # ── Build DOCX ─────────────────────────────────────────────────────────
    doc = Document()

    # — narrow margins —
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(80)
        section.right_margin = Pt(80)

    # Cover
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(script.topic_title)
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = COLOR_HDR

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run(
        f"PDF Podcast Agent  •  {datetime.now(timezone.utc).strftime('%d %B %Y')}\n"
        f"{len(script.lines)} voiced lines  •  {len(extended_lines)} extended lines"
    )
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # Episode Brief
    _set_heading(doc, "Episode Brief", level=1)

    _set_heading(doc, "Core Argument", level=2)
    doc.add_paragraph(insights.core_thesis).paragraph_format.space_after = Pt(8)

    _set_heading(doc, "Key Metrics", level=2)
    for m in insights.key_metrics:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(m).font.size = Pt(10.5)

    _set_heading(doc, "The Controversy", level=2)
    doc.add_paragraph(insights.controversial_point).paragraph_format.space_after = Pt(
        12
    )
    doc.add_page_break()

    # Voiced Script
    _set_heading(doc, "Podcast Script  (Voiced Version)", level=1)
    sub = doc.add_paragraph()
    sub.add_run(
        "These are the exact lines synthesised to audio by the Voice Studio agent."
    ).font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    ts = 0  # running timestamp in seconds (~12 s per TTS line)
    for line in script.lines:
        is_a = line.speaker == "HostA"
        stamp = f"[{ts // 60}:{ts % 60:02d}]"
        _add_speaker_block(
            doc,
            f"{stamp}  HOST A (The Skeptic)"
            if is_a
            else f"{stamp}  HOST B (The Expert)",
            line.text,
            COLOR_A if is_a else COLOR_B,
        )
        ts += 12
    doc.add_page_break()

    # Extended Script
    _set_heading(doc, "Extended Script  (Director's Cut)", level=1)
    sub2 = doc.add_paragraph()
    sub2.add_run(
        "Continues from the voiced script — deeper exchanges generated for the written version."
    ).font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    for line in extended_lines:
        spk = line.get("speaker", "HostA")
        txt = line.get("text", "")
        if not txt:
            continue
        is_a = spk == "HostA"
        _add_speaker_block(
            doc,
            "HOST A (The Skeptic)" if is_a else "HOST B (The Expert)",
            txt,
            COLOR_A if is_a else COLOR_B,
        )

    # Save alongside the MP3
    stem = Path(audio_path).stem
    docx_path = OUTPUT_DIR / f"{stem}.docx"
    doc.save(str(docx_path))
    return str(docx_path)


# ── Personality helpers ───────────────────────────────────────────────────────


def _load_personalities(ctx: Context) -> tuple[str, str, str, str]:
    """Return (a_name, a_hint, b_name, b_hint) from storage, falling back to defaults."""
    raw = ctx.storage.get(_PERSONALITIES_KEY)
    if raw:
        try:
            p = json.loads(raw)
            return p["a_name"], p["a_hint"], p["b_name"], p["b_hint"]
        except Exception:
            pass
    a_name, a_hint = _PERSONALITY_PRESETS["host_a"]["1"]
    b_name, b_hint = _PERSONALITY_PRESETS["host_b"]["1"]
    return a_name, a_hint, b_name, b_hint


# ── Payment helpers ───────────────────────────────────────────────────────────


def _stripe_enabled() -> bool:
    return bool(_stripe_lib and os.getenv("STRIPE_SECRET_KEY", ""))


def _is_session_paid(ctx: Context, session_id: str) -> bool:
    """Return True if this podcast session has an active Live Show Pass."""
    raw = ctx.storage.get(_PAID_SESSIONS_KEY) or "[]"
    try:
        return session_id in json.loads(raw)
    except Exception:
        return False


def _create_embedded_checkout(
    *, user_address: str, podcast_session_id: str, description: str
) -> dict:
    """Create a Stripe embedded Checkout Session and return the payload
    that goes into ``RequestPayment.metadata["stripe"]``.

    Uses ``ui_mode="embedded"`` so ASI:One renders the payment form
    inline as an overlay — no redirect required.
    """
    _stripe_lib.api_key = os.getenv("STRIPE_SECRET_KEY", "")  # type: ignore[union-attr]
    price_cents = int(os.getenv("STRIPE_LIVE_SHOW_PRICE_CENTS", "1000"))
    pub_key = os.getenv("STRIPE_PUBLISHABLE_KEY", "")
    success_url = os.getenv("STRIPE_SUCCESS_URL", "https://asi1.ai")

    session = _stripe_lib.checkout.Session.create(  # type: ignore[union-attr]
        ui_mode="embedded",
        redirect_on_completion="if_required",
        mode="payment",
        payment_method_types=["card"],
        return_url=success_url + "?session_id={CHECKOUT_SESSION_ID}",
        line_items=[
            {
                "price_data": {
                    "currency": "usd",
                    "unit_amount": price_cents,
                    "product_data": {
                        "name": "PDF Podcast - Live Show Pass",
                        "description": description,
                    },
                },
                "quantity": 1,
            }
        ],
        metadata={
            "podcast_session_id": podcast_session_id,
            "user_address": user_address,
        },
    )

    return {
        "client_secret": session.client_secret,
        "checkout_session_id": session.id,
        "publishable_key": pub_key,
        "currency": "usd",
        "amount_cents": price_cents,
        "ui_mode": "embedded_page",
    }


def _verify_checkout_session_paid(checkout_session_id: str) -> bool:
    """Verify with Stripe that a Checkout Session has been paid."""
    _stripe_lib.api_key = os.getenv("STRIPE_SECRET_KEY", "")  # type: ignore[union-attr]
    session = _stripe_lib.checkout.Session.retrieve(checkout_session_id)  # type: ignore[union-attr]
    return getattr(session, "payment_status", None) == "paid"


async def _gate_with_payment(ctx: Context, sender: str, feature_name: str) -> bool:
    """Check payment gate.  Returns True (proceed) or False (RequestPayment sent).

    Uses the Agent Payment Protocol with Stripe embedded Checkout so that
    ASI:One renders the payment form as an inline overlay in the chat.

    If STRIPE_SECRET_KEY is not set the gate is disabled and always returns True.
    """
    if not _stripe_enabled():
        ctx.logger.info("[Payment] Stripe not configured -- gate disabled (dev mode).")
        return True

    raw = ctx.storage.get(_LAST_SESSION_KEY)
    if not raw:
        await ctx.send(
            sender,
            _chat(
                "No podcast session found.\n\n"
                "Send me a PDF first to generate a podcast, then unlock the Live Show Pass."
            ),
        )
        return False

    data = json.loads(raw)
    podcast_session_id = data["session_id"]

    if _is_session_paid(ctx, podcast_session_id):
        return True

    price_cents = int(os.getenv("STRIPE_LIVE_SHOW_PRICE_CENTS", "1000"))
    price_dollars = price_cents / 100

    try:
        checkout_payload = await asyncio.to_thread(
            _create_embedded_checkout,
            user_address=sender,
            podcast_session_id=podcast_session_id,
            description=(
                "Unlock live debate, host personality customization "
                "& extended Q&A for this episode."
            ),
        )

        checkout_session_id = checkout_payload["checkout_session_id"]
        raw_pending = ctx.storage.get(_PENDING_PAYMENTS_KEY) or "{}"
        try:
            pending_by_checkout = json.loads(raw_pending)
            if not isinstance(pending_by_checkout, dict):
                pending_by_checkout = {}
        except Exception:
            pending_by_checkout = {}

        pending_by_checkout[checkout_session_id] = {
            "podcast_session_id": podcast_session_id,
            "user_address": sender,
            "feature": feature_name,
        }
        ctx.storage.set(_PENDING_PAYMENTS_KEY, json.dumps(pending_by_checkout))

        req = RequestPayment(
            accepted_funds=[
                Funds(
                    currency="USD",
                    amount=f"{price_dollars:.2f}",
                    payment_method="stripe",
                )
            ],
            recipient=str(ctx.agent.address),
            deadline_seconds=600,
            description=f"Pay ${price_dollars:.0f} to unlock the Live Show Pass for this episode.",
            metadata={"stripe": checkout_payload, "service": "live_show_pass"},
        )
        await ctx.send(sender, req)
        ctx.logger.info(
            f"[Payment] RequestPayment sent (embedded checkout "
            f"{checkout_payload['checkout_session_id'][:20]}...) "
            f"for session {podcast_session_id[:8]}"
        )
    except Exception as exc:
        ctx.logger.error(f"[Payment] Stripe error: {exc}")
        await ctx.send(
            sender,
            _chat(
                f"Payment service temporarily unavailable.\n\n"
                f"Please try again in a moment. *(Error: {exc})*"
            ),
        )
    return False


# ── Cover art ─────────────────────────────────────────────────────────────────

# ── Host-agent context injection ──────────────────────────────────────────────


async def _inject_context_to_hosts(
    ctx: Context,
    session_id: str,
    document_text: str,
    script: "PodcastScript",
    insights: "ResearchInsights",
) -> None:
    """Fire-and-forget context push to @HostA and @HostB so they can answer
    follow-up questions from ASI:One users after the podcast is ready."""
    if not HOST_A_ADDRESS and not HOST_B_ADDRESS:
        ctx.logger.warning(
            "[Orchestrator] HOST_A_ADDRESS / HOST_B_ADDRESS not set — skipping context injection"
        )
        return

    _, a_hint, _, b_hint = _load_personalities(ctx)

    payload = ContextInjection(
        session_id=session_id,
        topic_title=script.topic_title,
        core_thesis=insights.core_thesis,
        key_metrics=insights.key_metrics,
        controversial_point=insights.controversial_point,
        document_snippet=document_text[:4000],
        host_a_personality=a_hint,
        host_b_personality=b_hint,
    )

    for name, address in [("HostA", HOST_A_ADDRESS), ("HostB", HOST_B_ADDRESS)]:
        if address:
            await ctx.send(address, payload)
            ctx.logger.info(f"[Orchestrator] Context injected → {name}")


# ── Core pipeline ─────────────────────────────────────────────────────────────


async def _run_pipeline(
    ctx: Context, document_text: str, session_id: str
) -> tuple[AudioResponse, PodcastScript, ResearchInsights]:
    """
    Chain:  ExtractRequest → ResearchInsights → PodcastScript → AudioResponse
    Uses ctx.send_and_receive at each hop for clean request-response flow.
    """
    sid = session_id[:8]

    _PIPELINE_TIMEOUT = 120  # seconds — generous buffer for Agentverse relay

    # Step 1 — RAG Extraction
    ctx.logger.info(f"[{sid}] ⟶ Extractor …")
    insights, st1 = await ctx.send_and_receive(
        EXTRACTOR_ADDRESS,
        ExtractRequest(document_text=document_text, session_id=session_id),
        response_type=ResearchInsights,
        timeout=_PIPELINE_TIMEOUT,
    )
    if not isinstance(insights, ResearchInsights):
        raise RuntimeError(f"Extractor failed (status={st1})")
    ctx.logger.info(f"[{sid}] ✓ Insights — {insights.core_thesis[:60]}…")

    # Step 2 — Script generation
    ctx.logger.info(f"[{sid}] ⟶ Scriptwriter …")
    script, st2 = await ctx.send_and_receive(
        SCRIPTWRITER_ADDRESS,
        insights,
        response_type=PodcastScript,
        timeout=_PIPELINE_TIMEOUT,
    )
    if not isinstance(script, PodcastScript):
        raise RuntimeError(f"Scriptwriter failed (status={st2})")
    ctx.logger.info(
        f"[{sid}] ✓ Script — {len(script.lines)} lines  •  '{script.topic_title}'"
    )

    # Step 3 — Audio generation + stitching
    ctx.logger.info(f"[{sid}] ⟶ Voice Studio …")
    audio, st3 = await ctx.send_and_receive(
        VOICE_STUDIO_ADDRESS,
        script,
        response_type=AudioResponse,
        timeout=_PIPELINE_TIMEOUT,
    )
    if not isinstance(audio, AudioResponse):
        raise RuntimeError(f"Voice Studio failed (status={st3})")
    ctx.logger.info(f"[{sid}] ✓ Audio → {audio.audio_path}")

    return audio, script, insights


# ── Live debate (relay: HostA → Orchestrator → HostB → …, one growing bubble) ─

_MAX_DEBATE_TURNS = 8  # total lines (4 from each host)
_DEBATE_TURN_DELAY_SECS = int(os.getenv("DEBATE_TURN_DELAY_SECS", "8"))


async def _run_live_debate(ctx: Context, sender: str) -> None:
    """Kick off a turn-based debate.

    The orchestrator sends a DebateTurn to Host A, Host A replies with
    DebateResponse, the orchestrator appends it to the growing transcript and
    then (after a short pacing pause) sends the next DebateTurn to the other host.
    Every DebateResponse updates the *same* chat bubble with the full accumulated
    transcript so far — giving a "live-streaming" effect.
    """
    raw = ctx.storage.get(_LAST_SESSION_KEY)
    if not raw:
        await ctx.send(
            sender,
            _chat(
                "⚠️ No podcast session found yet.\n\n"
                "Send me a PDF first to generate a podcast, "
                "then say **debate** to watch the hosts go head-to-head."
            ),
        )
        return

    if not HOST_A_ADDRESS or not HOST_B_ADDRESS:
        await ctx.send(
            sender,
            _chat(
                "⚠️ Host agent addresses not configured.\n"
                "Set HOST_A_ADDRESS and HOST_B_ADDRESS in your .env file."
            ),
        )
        return

    data = json.loads(raw)
    topic = data["topic_title"]
    sid = str(uuid4())[:8]

    a_name, a_hint, b_name, b_hint = _load_personalities(ctx)

    # Stamp cooldown so duplicate trigger messages are ignored
    import time as _time

    now = str(_time.time())
    ctx.storage.set(_DEBATE_COOLDOWN_KEY, now)

    # Generate a stable msg_id for this debate — every bubble update reuses
    # it so ASI:One overwrites the same bubble instead of creating new ones.
    debate_msg_id = uuid4()
    ctx.storage.set(_DEBATE_MSG_ID_KEY, str(debate_msg_id))

    # Seed storage with empty transcript + active session + TTL timestamp
    ctx.storage.set(_ACTIVE_DEBATE_KEY, sid)
    ctx.storage.set(_ACTIVE_DEBATE_TS_KEY, now)
    ctx.storage.set(
        _DEBATE_ACCUM_KEY,
        json.dumps(
            {
                "session_id": sid,
                "lines": [],
                "user": sender,
                "a_hint": a_hint,
                "b_hint": b_hint,
            }
        ),
    )

    ctx.logger.info(
        f"[LiveDebate] Starting session {sid} for '{topic}' — A: {a_name}, B: {b_name}"
    )

    # Opening message — keep bubble OPEN so subsequent turns update it in-place
    header = (
        f"## 🎭 Live Debate — *{topic}*\n\n"
        f"*🎤 @skeptic-agent as **{a_name}** · 🎓 @expert-agent as **{b_name}***\n\n"
        f"*⏳ Warming up the hosts…*"
    )
    await ctx.send(sender, _chat_open(header, debate_msg_id))

    await ctx.send(
        HOST_A_ADDRESS,
        DebateTurn(
            session_id=sid,
            topic_title=data["topic_title"],
            core_thesis=data["core_thesis"],
            key_metrics=data.get("key_metrics", []),
            controversial_point=data.get("controversial_point", ""),
            document_snippet=data.get("document_snippet", "")[:2000],
            user_address=sender,
            turn=0,
            max_turns=_MAX_DEBATE_TURNS,
            previous_statement="",
            speaker_personality=a_hint,
        ),
    )
    ctx.logger.info(f"[LiveDebate] DebateTurn 0 → HostA ({a_name})")


# ── Agent ─────────────────────────────────────────────────────────────────────

_agentverse_key = os.getenv("AGENTVERSE_API_KEY", "")

orchestrator = Agent(
    name="pdf_podcast_orchestrator",
    seed=os.getenv("ORCHESTRATOR_SEED", "pdf_podcast_orchestrator_seed_v1"),
    port=8000,
    # Use mailbox ONLY so ASI:One can reach us via the Agentverse relay.
    # Do NOT set endpoint — it overrides mailbox in the Almanac, causing
    # ASI:One to try localhost:8000 (unreachable from the cloud).
    # Sub-agent replies route through Agentverse relay (slightly slower),
    # compensated by the 120 s timeout on send_and_receive calls.
    **(
        {
            "mailbox": _agentverse_key,
        }
        if _agentverse_key
        else {
            "endpoint": ["http://localhost:8000/submit"],
        }
    ),
    network="testnet",
)

chat_proto = Protocol(spec=chat_protocol_spec)

# ── Startup ───────────────────────────────────────────────────────────────────


@orchestrator.on_event("startup")
async def on_startup(ctx: Context) -> None:
    ctx.logger.info(f"[Orchestrator] address: {ctx.agent.address}")

    missing = [
        name
        for name, val in [
            ("EXTRACTOR_ADDRESS", EXTRACTOR_ADDRESS),
            ("SCRIPTWRITER_ADDRESS", SCRIPTWRITER_ADDRESS),
            ("VOICE_STUDIO_ADDRESS", VOICE_STUDIO_ADDRESS),
        ]
        if not val
    ]
    if missing:
        ctx.logger.warning(
            f"[Orchestrator] Missing env vars: {', '.join(missing)}\n"
            f"  Run `python get_addresses.py` to get the correct values,\n"
            f"  then set them before starting this agent."
        )
    else:
        ctx.logger.info(
            f"[Orchestrator] sub-agents wired:\n"
            f"  Extractor    {EXTRACTOR_ADDRESS}\n"
            f"  Scriptwriter {SCRIPTWRITER_ADDRESS}\n"
            f"  Voice Studio {VOICE_STUDIO_ADDRESS}\n"
            f"  Host A       {HOST_A_ADDRESS or '(not set)'}\n"
            f"  Host B       {HOST_B_ADDRESS or '(not set)'}"
        )
    ctx.storage.set(_PENDING_PAYMENTS_KEY, "{}")
    ctx.logger.info("[Payment] Pending payment state cleared on startup.")

    # Clear stale debate state — if the orchestrator restarted mid-debate the
    # active session key would block all future debate triggers indefinitely.
    ctx.storage.set(_ACTIVE_DEBATE_KEY, "")
    ctx.storage.set(_DEBATE_ACCUM_KEY, "")
    ctx.storage.set(_DEBATE_COOLDOWN_KEY, "0")
    ctx.storage.set(_ACTIVE_DEBATE_TS_KEY, "0")
    ctx.storage.set(_DEBATE_MSG_ID_KEY, "")
    ctx.storage.set(_SEEN_MSG_IDS_KEY, "[]")
    ctx.logger.info("[Debate] Stale debate state cleared on startup.")

    ctx.logger.info("[Orchestrator] REST: POST http://localhost:8000/process")


# ── Agent Payment Protocol (Stripe embedded checkout) ─────────────────────────

payment_proto = Protocol(spec=payment_protocol_spec, role="seller")


@payment_proto.on_message(CommitPayment)
async def on_commit_payment(ctx: Context, sender: str, msg: CommitPayment) -> None:
    """User completed the Stripe embedded checkout.

    ASI:One sends CommitPayment with ``transaction_id`` set to the Stripe
    Checkout Session ID.  We verify with Stripe, send CompletePayment,
    mark the session as paid, and notify the user.
    """
    if msg.funds.payment_method != "stripe" or not msg.transaction_id:
        await ctx.send(
            sender,
            RejectPayment(reason="Unsupported payment method (expected stripe)."),
        )
        return

    checkout_session_id = msg.transaction_id

    try:
        paid = await asyncio.to_thread(
            _verify_checkout_session_paid, checkout_session_id
        )
    except Exception as exc:
        ctx.logger.error(f"[Payment] Stripe verify error: {exc}")
        await ctx.send(
            sender,
            RejectPayment(
                reason="Could not verify payment with Stripe. Please try again."
            ),
        )
        return

    if not paid:
        await ctx.send(
            sender,
            RejectPayment(
                reason="Stripe payment not completed yet. Please finish checkout."
            ),
        )
        return

    await ctx.send(sender, CompletePayment(transaction_id=checkout_session_id))
    ctx.logger.info(f"[Payment] CompletePayment sent for {checkout_session_id[:20]}...")

    # Determine which podcast session this checkout belongs to.
    # Pending payments are keyed by checkout_session_id to avoid races between
    # concurrent users paying at the same time.
    raw_pending = ctx.storage.get(_PENDING_PAYMENTS_KEY) or "{}"
    try:
        pending_data = json.loads(raw_pending)
        if not isinstance(pending_data, dict):
            pending_data = {}
    except Exception:
        pending_data = {}

    pending_info = pending_data.get(checkout_session_id)
    if not pending_info and {"checkout_session_id", "podcast_session_id"}.issubset(
        pending_data.keys()
    ):
        # Backward compatibility for older single-pending-payment shape.
        legacy_checkout = pending_data.get("checkout_session_id")
        if legacy_checkout == checkout_session_id:
            pending_info = pending_data

    podcast_session_id = ""
    if isinstance(pending_info, dict):
        podcast_session_id = pending_info.get("podcast_session_id", "")
    if not isinstance(podcast_session_id, str):
        podcast_session_id = ""

    if not podcast_session_id:
        # Fallback: retrieve from Stripe metadata
        try:
            _stripe_lib.api_key = os.getenv("STRIPE_SECRET_KEY", "")  # type: ignore[union-attr]
            stripe_session = await asyncio.to_thread(
                _stripe_lib.checkout.Session.retrieve,
                checkout_session_id,  # type: ignore[union-attr]
            )
            podcast_session_id = (stripe_session.metadata or {}).get(
                "podcast_session_id", ""
            )
        except Exception:
            pass

    if podcast_session_id:
        raw_paid = ctx.storage.get(_PAID_SESSIONS_KEY) or "[]"
        paid_ids: list = json.loads(raw_paid)
        if podcast_session_id not in paid_ids:
            paid_ids.append(podcast_session_id)
        ctx.storage.set(_PAID_SESSIONS_KEY, json.dumps(paid_ids))
        ctx.logger.info(f"[Payment] Session {podcast_session_id[:8]} marked as paid.")

    if checkout_session_id in pending_data:
        pending_data.pop(checkout_session_id, None)
        ctx.storage.set(_PENDING_PAYMENTS_KEY, json.dumps(pending_data))

    await ctx.send(
        sender,
        _chat(
            "## Payment confirmed -- Live Show Pass is active!\n\n"
            "The hosts are loaded and ready. Here's what you can do now:\n\n"
            "* **Start the live debate** -- type `continue debate`\n"
            "* **Customize host personalities first** -- type `A:[1-4] B:[1-4]` "
            "*(e.g. `A:2 B:3`)* or type `customize` to see the full menu\n"
            "* **Ask the hosts anything** -- tag @skeptic-agent or @expert-agent\n\n"
            "*Tip: set personalities before the debate for the best experience!*"
        ),
    )


@payment_proto.on_message(RejectPayment)
async def on_reject_payment(ctx: Context, sender: str, msg: RejectPayment) -> None:
    """User cancelled or the UI rejected the payment."""
    ctx.logger.info(f"[Payment] Payment rejected by {sender[:20]}...: {msg.reason}")


# ── Debate response relay (Host → Orchestrator → User bubble) ─────────────────


@orchestrator.on_message(DebateResponse)
async def handle_debate_response(
    ctx: Context, sender: str, msg: DebateResponse
) -> None:
    """Receive one debate line, accumulate internally, show progress only.

    Intermediate updates send a short progress indicator (no debate content)
    so the chat stays clean.  The complete transcript is delivered in a single
    message only after the final turn.
    """
    active_sid = ctx.storage.get(_ACTIVE_DEBATE_KEY)
    if active_sid != msg.session_id:
        ctx.logger.warning(
            f"[DebateResponse] Stale session {msg.session_id[:8]} (active={active_sid}). Dropped."
        )
        return

    raw_accum = ctx.storage.get(_DEBATE_ACCUM_KEY)
    if not raw_accum:
        ctx.logger.error("[DebateResponse] No accumulator found. Abort.")
        return
    accum = json.loads(raw_accum)

    debate_msg_id = UUID(ctx.storage.get(_DEBATE_MSG_ID_KEY) or str(uuid4()))

    label = (
        "🎤 **@skeptic-agent**" if msg.speaker == "skeptic" else "🎓 **@expert-agent**"
    )
    accum["lines"].append(f"{label}\n{msg.reply_text}")
    ctx.storage.set(_DEBATE_ACCUM_KEY, json.dumps(accum))

    user_address = accum["user"]
    a_hint = accum.get("a_hint", "")
    b_hint = accum.get("b_hint", "")
    topic = msg.topic_title
    total = len(accum["lines"])
    next_turn = msg.turn + 1
    is_final = next_turn >= msg.max_turns

    if is_final:
        # Deliver the COMPLETE debate transcript in one message
        body = "\n\n".join(accum["lines"])
        full_text = (
            f"## 🎭 Live Debate — *{topic}*\n\n"
            f"{body}\n\n"
            "---\n"
            "🎬 *Debate over. Tag @skeptic-agent or @expert-agent to keep the conversation going.*"
        )
        await ctx.send(user_address, _chat(full_text))
        ctx.storage.set(_ACTIVE_DEBATE_KEY, "")
        ctx.logger.info(
            f"[DebateResponse] Debate finished — {msg.max_turns} turns delivered."
        )
    else:
        next_speaker = "expert" if msg.speaker == "skeptic" else "skeptic"
        next_label = (
            "🎓 @expert-agent" if next_speaker == "expert" else "🎤 @skeptic-agent"
        )
        next_personality = b_hint if next_speaker == "expert" else a_hint

        # Progress-only update — no debate content, keeps the chat clean
        speaker_done = (
            "🎤 @skeptic-agent" if msg.speaker == "skeptic" else "🎓 @expert-agent"
        )
        progress_bar = "".join("█" if i < total else "░" for i in range(msg.max_turns))
        progress_text = (
            f"## 🎭 Live Debate — *{topic}*\n\n"
            f"`{progress_bar}` {total}/{msg.max_turns}\n\n"
            f"✓ {speaker_done} — *delivered*\n\n"
            f"⏳ *{next_label} is responding…*"
        )
        await ctx.send(user_address, _chat_open(progress_text, debate_msg_id))
        ctx.logger.info(
            f"[DebateResponse] Turn {msg.turn} delivered. Waiting {_DEBATE_TURN_DELAY_SECS} s for turn {next_turn}…"
        )

        # Keep a human-paced cadence in the "live" stream while preserving turn order.
        await asyncio.sleep(_DEBATE_TURN_DELAY_SECS)

        # Build a plain-text debate history from all accumulated lines
        history_lines = []
        for line in accum["lines"]:
            clean = line.replace("🎤 **@skeptic-agent**\n", "Skeptic: ")
            clean = clean.replace("🎓 **@expert-agent**\n", "Expert: ")
            history_lines.append(clean)
        debate_history = "\n".join(history_lines)

        next_host = HOST_B_ADDRESS if next_speaker == "expert" else HOST_A_ADDRESS
        await ctx.send(
            next_host,
            DebateTurn(
                session_id=msg.session_id,
                topic_title=msg.topic_title,
                core_thesis=msg.core_thesis,
                key_metrics=msg.key_metrics,
                controversial_point=msg.controversial_point,
                document_snippet=msg.document_snippet,
                user_address=user_address,
                turn=next_turn,
                max_turns=msg.max_turns,
                previous_statement=msg.reply_text,
                speaker_personality=next_personality,
                debate_history=debate_history,
            ),
        )
        ctx.logger.info(
            f"[DebateResponse] DebateTurn {next_turn} → "
            f"{'HostB' if next_speaker == 'expert' else 'HostA'}"
        )


# ── ASI:One Chat Protocol ─────────────────────────────────────────────────────


@chat_proto.on_message(ChatMessage)
async def handle_chat_message(ctx: Context, sender: str, msg: ChatMessage) -> None:
    await ctx.send(
        sender,
        ChatAcknowledgement(
            timestamp=datetime.now(timezone.utc), acknowledged_msg_id=msg.msg_id
        ),
    )

    # ── Message deduplication (Agentverse retries the same msg_id) ────────────
    import time as _time

    msg_key = str(msg.msg_id)
    try:
        seen_raw = ctx.storage.get(_SEEN_MSG_IDS_KEY) or "[]"
        seen: list = json.loads(seen_raw)
    except Exception:
        seen = []
    if msg_key in seen:
        ctx.logger.info(f"[Chat] Duplicate msg_id {msg_key[:12]}… — skipped.")
        return
    seen.append(msg_key)
    seen = seen[-50:]  # keep last 50 to avoid unbounded growth
    ctx.storage.set(_SEEN_MSG_IDS_KEY, json.dumps(seen))

    # ── Auto-expire stale active debates (TTL guard) ─────────────────────────
    active = ctx.storage.get(_ACTIVE_DEBATE_KEY)
    if active:
        debate_ts = float(ctx.storage.get(_ACTIVE_DEBATE_TS_KEY) or "0")
        if _time.time() - debate_ts > _ACTIVE_DEBATE_TTL:
            ctx.logger.info(
                f"[Chat] Stale debate {active[:8]} expired after "
                f"{_ACTIVE_DEBATE_TTL}s — clearing."
            )
            ctx.storage.set(_ACTIVE_DEBATE_KEY, "")
            ctx.storage.set(_DEBATE_ACCUM_KEY, "")

    # ── Text-command block (debate, personality) ──────────────────────────────
    has_resource = any(isinstance(item, ResourceContent) for item in msg.content)
    if not has_resource:
        combined_text = " ".join(
            item.text.strip()
            for item in msg.content
            if isinstance(item, TextContent) and item.text.strip()
        )
        combined_lower = combined_text.lower()

        # Personality selection: "A:2 B:3"
        personality_match = _PERSONALITY_SET_RE.search(combined_text)
        if personality_match:
            if not await _gate_with_payment(
                ctx, sender, "Host Personality Customization"
            ):
                return
            a_key, b_key = personality_match.group(1), personality_match.group(2)
            a_name, a_hint = _PERSONALITY_PRESETS["host_a"][a_key]
            b_name, b_hint = _PERSONALITY_PRESETS["host_b"][b_key]
            ctx.storage.set(
                _PERSONALITIES_KEY,
                json.dumps(
                    {
                        "a": a_key,
                        "b": b_key,
                        "a_name": a_name,
                        "a_hint": a_hint,
                        "b_name": b_name,
                        "b_hint": b_hint,
                    }
                ),
            )
            ctx.logger.info(f"[Personalities] Set A={a_name}, B={b_name}")
            await ctx.send(
                sender,
                _chat(
                    f"✅ **Host personalities updated!**\n\n"
                    f"🎤 **@skeptic-agent** → *{a_name}*\n"
                    f"🎓 **@expert-agent** → *{b_name}*\n\n"
                    f"These apply to all Q&A responses and the next debate.\n"
                    f"Type `debate` to watch them go head-to-head with the new styles."
                ),
            )
            return

        # Personality picker menu
        if any(t in combined_lower for t in _PERSONALITY_TRIGGERS):
            if not await _gate_with_payment(
                ctx, sender, "Host Personality Customization"
            ):
                return
            await ctx.send(sender, _chat(_PERSONALITY_MENU))
            return

        # Live debate trigger
        if any(trigger in combined_lower for trigger in _LIVE_DEBATE_TRIGGERS):
            # Cooldown guard — ASI:One retries cause duplicate deliveries
            last_ts = float(ctx.storage.get(_DEBATE_COOLDOWN_KEY) or "0")
            if _time.time() - last_ts < _DEBATE_COOLDOWN_SECS:
                ctx.logger.info(
                    f"[Chat] Debate trigger debounced (cooldown {_DEBATE_COOLDOWN_SECS}s active)."
                )
                return
            # Active-debate guard — don't spawn a second debate while one is running
            active = ctx.storage.get(_ACTIVE_DEBATE_KEY)
            if active:
                ctx.logger.info(
                    f"[Chat] Debate trigger ignored — session {active[:8]} already active."
                )
                return
            if not await _gate_with_payment(ctx, sender, "Live Debate"):
                return
            await _run_live_debate(ctx, sender)
            return

    document_text: Optional[str] = None

    for item in msg.content:
        # PDF sent as an attachment from ASI:One
        if isinstance(item, ResourceContent):
            if isinstance(item.resource, list) and not item.resource:
                ctx.logger.warning(
                    "[Chat] ResourceContent had empty resource list; skipping item."
                )
                continue
            resource = (
                item.resource[0] if isinstance(item.resource, list) else item.resource
            )
            try:
                document_text = _resource_to_text(resource)
                ctx.logger.info(
                    f"[Chat] PDF attachment received via ResourceContent ({len(document_text):,} chars)"
                )
            except Exception as e:
                ctx.logger.error(f"ResourceContent PDF read error: {e}")
            break

        if isinstance(item, TextContent):
            text = item.text.strip()
            if len(text) > 300 and not text.lower().endswith(".pdf"):
                document_text = text
                break
            if text.lower().endswith(".pdf") and os.path.exists(text):
                try:
                    document_text = _pdf_path_to_text(text)
                except Exception as e:
                    ctx.logger.error(f"PDF read error: {e}")
                break

    if not document_text:
        await ctx.send(
            sender,
            _chat(
                "Hi! I'm the PDF Podcast Agent.\n\n"
                "Send me either:\n"
                "  • The full pasted text of your PDF, or\n"
                "  • An absolute server-side path to a .pdf file\n\n"
                "I'll turn it into a 2-host debate podcast!"
            ),
        )
        return

    session_id = str(uuid4())

    await ctx.send(
        sender,
        _chat(
            f"Starting pipeline — session `{session_id[:8]}`\n"
            f"Steps: Extractor → Scriptwriter → Voice Studio\n"
            f"Usually 30–90 s. Hang tight!"
        ),
    )

    try:
        audio, script, insights = await _run_pipeline(ctx, document_text, session_id)

        # Persist session context so the live-debate trigger can use it later
        ctx.storage.set(
            _LAST_SESSION_KEY,
            json.dumps(
                {
                    "session_id": session_id,
                    "topic_title": script.topic_title,
                    "core_thesis": insights.core_thesis,
                    "key_metrics": insights.key_metrics,
                    "controversial_point": insights.controversial_point,
                    "document_snippet": document_text[:3000],
                    "script_lines": [
                        {"speaker": line.speaker, "text": line.text}
                        for line in script.lines
                    ],
                }
            ),
        )

        audio_url = _audio_url(audio.audio_path)
        line_count = len(script.lines)
        duration_s = line_count * 12
        duration = f"{duration_s // 60}:{duration_s % 60:02d}"

        # Run post-pipeline tasks concurrently
        docx_path, _ = await asyncio.gather(
            _build_docx(audio.audio_path, script, insights),
            _inject_context_to_hosts(ctx, session_id, document_text, script, insights),
        )

        docx_url = _audio_url(docx_path)

        # ── Consolidated episode snapshot ───────────────────────────────────
        snapshot_lines = [
            f"> 📖 **{insights.core_thesis[:120]}{'…' if len(insights.core_thesis) > 120 else ''}**",
            ">",
        ]
        for m in insights.key_metrics:
            snapshot_lines.append(f"> • {m}")
        snapshot_lines += [
            ">",
            f"> 🔥 *{insights.controversial_point[:120]}{'…' if len(insights.controversial_point) > 120 else ''}*",
            ">",
        ]
        for line in script.lines[:3]:
            tag = "[A]" if line.speaker == "HostA" else "[B]"
            snapshot_lines.append(
                f"> **{tag}** {line.text[:90]}{'…' if len(line.text) > 90 else ''}"
            )
        if line_count > 3:
            snapshot_lines.append("> …")
        snapshot = "\n".join(snapshot_lines)

        # ── Live Q&A section ─────────────────────────────────────────────────
        hosts_ready = bool(HOST_A_ADDRESS and HOST_B_ADDRESS)
        session_paid = _is_session_paid(ctx, session_id)
        stripe_active = _stripe_enabled()
        a_name, _, b_name, _ = _load_personalities(ctx)
        price_dollars = int(os.getenv("STRIPE_LIVE_SHOW_PRICE_CENTS", "1000")) / 100

        if not hosts_ready:
            qa_section = (
                "## 💬 Live Q&A — Ask the Hosts\n"
                "*(Start `host_a_agent.py` and `host_b_agent.py` and add their addresses "
                "to `.env` as `HOST_A_ADDRESS` / `HOST_B_ADDRESS` to enable live Q&A and debate.)*"
            )
        elif stripe_active and not session_paid:
            # Payment required — show free tier + locked premium section
            qa_section = (
                "## 💬 Live Q&A — Ask the Hosts *(FREE)*\n"
                "I've called **@skeptic-agent** and **@expert-agent** into this chat "
                "and they have memorised the full paper.\n\n"
                "Use **@skeptic-agent** or **@expert-agent** here to ask anything about the episode.\n\n"
                "---\n\n"
                f"## 🔒 Live Show Pass — ${price_dollars:.0f}\n"
                "Unlock the full interactive experience for this episode:\n\n"
                "• 🎭 **Live Debate** — @skeptic-agent vs @expert-agent, streamed turn-by-turn\n"
                "• 🎨 **Host Personality Customization** — 16 personality combos\n"
                "• 💬 **Extended Q&A** — deeper, richer host responses\n\n"
                "Type **`debate`** or **`customize`** to unlock via Stripe."
            )
        else:
            # No payment gate (Stripe not configured) OR already paid — show everything
            pass_badge = " *(🎟️ Live Show Pass active)*" if session_paid else ""
            qa_section = (
                f"## 💬 Live Q&A — Ask the Hosts{pass_badge}\n"
                "I've called **@skeptic-agent** and **@expert-agent** into this chat "
                "and they have memorised the full paper.\n\n"
                "Use **@skeptic-agent** or **@expert-agent** here to ask anything about the episode.\n\n"
                "---\n\n"
                "## 🎭 Watch Them Debate Live\n"
                "Want to see @skeptic-agent and @expert-agent go head-to-head on this paper?\n\n"
                "Just type any of these:\n"
                "> `debate` &nbsp;·&nbsp; `live debate` &nbsp;·&nbsp; `start debate` "
                "&nbsp;·&nbsp; `watch them` &nbsp;·&nbsp; `replay`\n\n"
                "---\n\n"
                "## 🎨 Customize Host Personalities\n"
                f"*Currently: 🎤 **{a_name}** · 🎓 **{b_name}***\n\n"
                "Want a different vibe? Reply **A:[1-4] B:[1-4]**:\n\n"
                "🎤 **@skeptic-agent**: 1. Classic Skeptic &nbsp;·&nbsp; 2. Investigative Journalist "
                "&nbsp;·&nbsp; 3. Academic Critic &nbsp;·&nbsp; 4. Industry Veteran\n\n"
                "🎓 **@expert-agent**: 1. Researcher &nbsp;·&nbsp; 2. Industry Insider "
                "&nbsp;·&nbsp; 3. Futurist &nbsp;·&nbsp; 4. Enthusiastic Teacher\n\n"
                "*Example: type* `A:2 B:3` *for Investigative Journalist vs Futurist.*"
            )

        reply = "\n".join(
            [
                "## ✅ Podcast Ready!",
                "",
                f"## 🎙️ {script.topic_title}",
                f"⏱️ ~{duration} &nbsp;•&nbsp; {line_count} voiced lines &nbsp;•&nbsp; 2 hosts",
                "",
                "---",
                "## 🔗 Downloads",
                f"🔊 **Listen (MP3):** {audio_url}",
                f"📄 **Full Script (DOCX):** {docx_url}",
                "",
                "---",
                "## 📋 Episode Snapshot",
                "*Preview only — open the DOCX for the full analysis, "
                "timestamped transcript, and extended director's cut.*",
                "",
                snapshot,
                "",
                "---",
                # ── Change 3: single-line background note ──────────────────────
                "🧠 *Paper context loaded into host agents.*",
                "",
                "---",
                qa_section,
            ]
        )

        await ctx.send(sender, _chat(reply))

    except Exception as exc:
        ctx.logger.error(f"Pipeline error: {exc}")
        await ctx.send(sender, _chat(f"❌ Pipeline failed: {exc}"))


@chat_proto.on_message(ChatAcknowledgement)
async def handle_ack(ctx: Context, sender: str, msg: ChatAcknowledgement) -> None:
    ctx.logger.info(f"[Orchestrator] ACK received from {sender[:20]}…")


# ── REST endpoint (local testing) ─────────────────────────────────────────────


@orchestrator.on_rest_post("/process", PipelineRequest, PipelineResponse)
async def handle_rest_process(ctx: Context, req: PipelineRequest) -> PipelineResponse:
    session_id = str(uuid4())
    ctx.logger.info(f"[REST] session {session_id[:8]}")

    try:
        if req.pdf_base64:
            document_text = _pdf_bytes_to_text(base64.b64decode(req.pdf_base64))
        elif req.pdf_path:
            document_text = _pdf_path_to_text(req.pdf_path)
        else:
            return PipelineResponse(
                audio_base64="",
                audio_path="",
                script_json="[]",
                topic_title="",
                status="error",
                error_message="Provide pdf_path or pdf_base64.",
            )

        if not document_text.strip():
            return PipelineResponse(
                audio_base64="",
                audio_path="",
                script_json="[]",
                topic_title="",
                status="error",
                error_message="Could not extract text from PDF.",
            )

        ctx.logger.info(
            f"[REST] {len(document_text):,} chars extracted — running pipeline …"
        )
        audio, script, insights = await _run_pipeline(ctx, document_text, session_id)

        # Voice Studio no longer sends base64 over the wire to stay under the
        # Agentverse mailbox size limit — read the file from disk instead.
        audio_b64 = audio.audio_base64
        if not audio_b64 and audio.audio_path:
            try:
                audio_b64 = base64.b64encode(
                    Path(audio.audio_path).read_bytes()
                ).decode()
            except Exception:
                audio_b64 = ""

        return PipelineResponse(
            audio_base64=audio_b64,
            audio_path=audio.audio_path,
            script_json=json.dumps(
                [{"speaker": line.speaker, "text": line.text} for line in script.lines],
                indent=2,
            ),
            topic_title=script.topic_title,
            status="success",
        )

    except Exception as exc:
        ctx.logger.error(f"[REST] error: {exc}")
        return PipelineResponse(
            audio_base64="",
            audio_path="",
            script_json="[]",
            topic_title="",
            status="error",
            error_message=str(exc),
        )


orchestrator.include(chat_proto, publish_manifest=True)
orchestrator.include(payment_proto, publish_manifest=True)


if __name__ == "__main__":
    orchestrator.run()
